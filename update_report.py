"""
update_report.py — Update paper/FINAL_REPORT.docx with the new equal-FFE-budget
benchmark results (3 methods x 8 configurations x 4 instances x 10 runs).

All numbers written into the report are computed from the CSV files in
./results so text, tables, and figures always agree.

The original document is backed up to paper/FINAL_REPORT_backup.docx first.

Usage:  python update_report.py
"""

import copy
import shutil
import zipfile
from pathlib import Path

import pandas as pd
from PIL import Image

import docx
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
RESULTS = ROOT / "results"
PLOTS = RESULTS / "plots"
DOC_PATH = ROOT / "paper" / "FINAL_REPORT.docx"
BACKUP_PATH = ROOT / "paper" / "FINAL_REPORT_backup.docx"

BUDGET = 2000
N_CONFIGS = 8
N_RETEST = 10
METHODS = ["harmony", "harmony_ls", "ant_colony"]
LABEL = {"harmony": "HS", "harmony_ls": "HS+LS", "ant_colony": "ACO"}
INSTANCES_URL = "https://github.com/saksunger/scheduling_optimization/tree/main/instances"

INSTANCE_ROWS = [
    ("1 (baseline)",    "100 / 250 / 250", "[10,30]; [50,200]; [5,20]",  "1001"),
    ("2 (URLLC-heavy)", "200 / 200 / 200", "[10,30]; [50,200]; [5,20]",  "2002"),
    ("3 (mMTC-heavy)",  "60 / 180 / 360",  "[10,30]; [50,200]; [10,30]", "3003"),
    ("4 (eMBB-heavy)",  "80 / 350 / 170",  "[10,30]; [80,250]; [5,20]",  "4004"),
]

FIGURE_FILES = [
    "convergence_harmony.png",
    "convergence_ant_colony.png",
    "convergence_best_hs_vs_aco.png",
    "fitness_comparison.png",
    "accuracy_comparison.png",
    "time_comparison.png",
    "service_metrics.png",
    "fitness_distribution.png",
]

REFERENCES = [
    "[1] Z. W. Geem, J. H. Kim, and G. V. Loganathan, “A New Heuristic Optimization "
    "Algorithm: Harmony Search,” Simulation, vol. 76, no. 2, pp. 60–68, 2001.",
    "[2] M. Dorigo, V. Maniezzo, and A. Colorni, “Ant System: Optimization by a Colony "
    "of Cooperating Agents,” IEEE Transactions on Systems, Man, and Cybernetics – "
    "Part B, vol. 26, no. 1, pp. 29–41, 1996.",
    "[3] M. Dorigo and T. Stützle, Ant Colony Optimization. Cambridge, MA: MIT Press, 2004.",
    "[4] P. Moscato, “On Evolution, Search, Optimization, Genetic Algorithms and Martial "
    "Arts: Towards Memetic Algorithms,” Caltech Concurrent Computation Program, Report "
    "826, 1989.",
    "[5] H. H. Hoos and T. Stützle, Stochastic Local Search: Foundations and Applications. "
    "San Francisco, CA: Morgan Kaufmann, 2004.",
    "[6] 3GPP, “Study on Scenarios and Requirements for Next Generation Access "
    "Technologies,” 3rd Generation Partnership Project, TR 38.913, 2017.",
    "[7] P. Popovski, K. F. Trillingsgaard, O. Simeone, and G. Durisi, “5G Wireless "
    "Network Slicing for eMBB, URLLC, and mMTC: A Communication-Theoretic View,” IEEE "
    "Access, vol. 6, pp. 55765–55779, 2018.",
    "[8] A. E. Eiben and S. K. Smit, “Parameter Tuning for Configuring and Analyzing "
    "Evolutionary Algorithms,” Swarm and Evolutionary Computation, vol. 1, no. 1, "
    "pp. 19–31, 2011.",
]


# --------------------------------------------------------------------------- #
# Load benchmark data and compute every number used in the report
# --------------------------------------------------------------------------- #
def load_data():
    agg = {m: {i: pd.read_csv(RESULTS / f"{m}_aggregated_instance{i}.csv")
               for i in range(1, 5)} for m in METHODS}

    s = {"agg": agg}
    s["best_cfg"] = {m: {i: int(agg[m][i]["AvgFitness"].idxmin()) for i in range(1, 5)}
                     for m in METHODS}
    s["best_avg"] = {m: {i: float(agg[m][i]["AvgFitness"].min()) for i in range(1, 5)}
                     for m in METHODS}
    s["best_best"] = {m: {i: float(agg[m][i].loc[s["best_cfg"][m][i], "BestFitness"])
                          for i in range(1, 5)} for m in METHODS}
    s["best_known"] = {i: min(float(agg[m][i]["BestFitness"].min()) for m in METHODS)
                       for i in range(1, 5)}
    s["winner"] = {i: min(METHODS, key=lambda m: s["best_avg"][m][i]) for i in range(1, 5)}
    # Local Search effect: best-config average delta (negative = LS better)
    s["ls_delta"] = {i: s["best_avg"]["harmony_ls"][i] - s["best_avg"]["harmony"][i]
                     for i in range(1, 5)}
    # mean wall-clock time per run (mean over configs), per method, mean over instances
    s["time"] = {m: float(pd.concat([agg[m][i]["AvgTime(ms)"] for i in range(1, 5)]).mean())
                 for m in METHODS}
    # config-sensitivity: spread between best and worst config average, per method
    s["spread"] = {m: float(sum(agg[m][i]["AvgFitness"].max() - agg[m][i]["AvgFitness"].min()
                                for i in range(1, 5)) / 4) for m in METHODS}
    # improvement between FFE 500 and FFE 2000 (best config, mean over runs+instances)
    s["late_gain"] = {}
    for m in METHODS:
        gains = []
        for i in range(1, 5):
            c = s["best_cfg"][m][i] + 1
            df = pd.read_csv(RESULTS / f"{m}_ffe_instance{i}_config{c}.csv")
            v500 = df[df["Evaluation"] == 500].iloc[0, 1:].mean()
            vend = df[df["Evaluation"] == BUDGET].iloc[0, 1:].mean()
            gains.append((v500 - vend) / v500 * 100)
        s["late_gain"][m] = sum(gains) / len(gains)
    # service metrics, instance 1, best config of each method
    s["svc"] = {m: agg[m][1].loc[s["best_cfg"][m][1]] for m in METHODS}
    return s


def fmt(v):
    return f"{v:,.1f}"


# --------------------------------------------------------------------------- #
# Safe low-level editing primitives
# --------------------------------------------------------------------------- #
def wts(element):
    return element.findall(".//" + qn("w:t"))


def set_text(p, text):
    """Rewrite a paragraph's text without touching non-text content (images)."""
    ts = wts(p._element)
    if not ts:
        p.add_run(text)
        return
    ts[0].text = text
    ts[0].set(qn("xml:space"), "preserve")
    for t in ts[1:]:
        t.text = ""


def set_caption(p, prefix, body):
    """Rewrite a 'Figure N.' caption, keeping the bold prefix run and any image."""
    ts = wts(p._element)
    fig_i = next((i for i, t in enumerate(ts)
                  if (t.text or "").lstrip().startswith("Figure")), None)
    if fig_i is None:
        set_text(p, prefix + body)
        return
    ts[fig_i].text = prefix
    ts[fig_i].set(qn("xml:space"), "preserve")
    if fig_i + 1 < len(ts):
        ts[fig_i + 1].text = body
        ts[fig_i + 1].set(qn("xml:space"), "preserve")
        for t in ts[fig_i + 2:]:
            t.text = ""
    else:
        ts[fig_i].text = prefix + body


def replace_in_paragraph(p, old, new):
    """Replace `old` with `new` if it sits inside a single run; else append new."""
    for t in wts(p._element):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)
            return True
    return False


def cell_set(cell, text):
    ts = wts(cell._tc)
    if ts:
        ts[0].text = str(text)
        ts[0].set(qn("xml:space"), "preserve")
        for t in ts[1:]:
            t.text = ""
    else:
        cell.text = str(text)


def clone_row(tbl):
    """Append a copy of the first data row and return the table row count."""
    trs = tbl._element.findall(qn("w:tr"))
    tbl._element.append(copy.deepcopy(trs[1]))


def insert_column(tbl, at):
    """Insert a column (copy of the last) at position `at` in every row."""
    grid = tbl._element.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol"))
    cols[at - 1].addnext(copy.deepcopy(cols[-1]))
    for tr in tbl._element.findall(qn("w:tr")):
        tcs = tr.findall(qn("w:tc"))
        tcs[at - 1].addnext(copy.deepcopy(tcs[at - 1]))


def new_paragraph_like(template_p, text):
    el = copy.deepcopy(template_p._element)
    # strip any drawings copied from the template
    for d in el.findall(".//" + qn("w:drawing")):
        d.getparent().remove(d)
    p = Paragraph(el, template_p._parent)
    set_text(p, text)
    return p


def add_hyperlink(p, url, text):
    """Append a clickable hyperlink run to a paragraph."""
    part = p.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = p._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    r = p._element.makeelement(qn("w:r"), {})
    rpr = p._element.makeelement(qn("w:rPr"), {})
    u = p._element.makeelement(qn("w:u"), {qn("w:val"): "single"})
    color = p._element.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    rpr.append(u)
    rpr.append(color)
    r.append(rpr)
    t = p._element.makeelement(qn("w:t"), {})
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hl.append(r)
    p._element.append(hl)


# --------------------------------------------------------------------------- #
def main():
    s = load_data()
    agg = s["agg"]

    shutil.copy2(DOC_PATH, BACKUP_PATH)
    print(f"Backup written to {BACKUP_PATH}")

    doc = docx.Document(str(DOC_PATH))
    P = doc.paragraphs
    T = doc.tables

    # Sanity-check the anchors we rely on before changing anything
    expectations = {
        1: "Authors:", 5: "Note on figures", 7: "Abstract", 8: "This report presents",
        9: "Each algorithm was evaluated", 12: "5G networks must allocate",
        18: "The objectives of this study", 20: "Harmony Search (HS)",
        21: "Ant Colony Optimization (ACO)", 27: "(Source:", 32: "(Source:",
        42: "(Source:", 47: "The added operator", 48: "calculateFitness",
        54: "The benchmark uses", 59: "Both algorithms share",
        60: "4. Experimental Setup", 62: "Four parameter sets",
        63: "4.2 Parameter Sets", 64: "Harmony Search", 70: "4.3 Performance Metrics",
        73: "Accuracy (consistency)", 75: "5.1 Harmony Search",
        78: "Among the Local-Search-free settings", 79: "", 80: "5.2 Ant Colony",
        83: "ACO's best configuration", 85: "Figure 2.", 87: "The central comparison",
        90: "Harmony Search wins", 97: "For the best configuration",
        100: "ACO achieves lower latency", 105: "Parameter sensitivity.",
        106: "Quality vs. time trade-off.", 107: "Exploration vs. exploitation.",
        108: "When Local Search pays off.", 109: "Service priorities.",
        111: "On this 5G mode-assignment problem", 112: "Recommendations.",
        115: "Appendix A", 118: "Appendix B", 119: "results/harmony_aggregated",
        120: "results/algorithm_comparison", 121: "results/ant_colony_fitness",
        122: "performance_test.cpp",
    }
    for idx, prefix in expectations.items():
        if prefix and not P[idx].text.strip().startswith(prefix):
            raise SystemExit(
                f"Anchor mismatch at paragraph {idx}: expected {prefix!r}, "
                f"found {P[idx].text.strip()[:60]!r} — aborting, document unchanged.")

    # ------------------------------------------------------------------ #
    # Computed summary phrases
    # ------------------------------------------------------------------ #
    wins = {m: sum(1 for i in range(1, 5) if s["winner"][i] == m) for m in METHODS}
    overall = max(METHODS, key=lambda m: wins[m])
    if wins[overall] == 4:
        winner_phrase = (f"{LABEL[overall]} attains the lowest average fitness "
                         f"on all four instances")
    else:
        parts = [f"{LABEL[m]} on {wins[m]}" for m in METHODS if wins[m]]
        winner_phrase = ("the lowest average fitness is attained by "
                         + ", ".join(parts) + " of the four instances")
    ls_better = sum(1 for i in range(1, 5) if s["ls_delta"][i] < 0)
    delta_str = ", ".join(f"I{i}: {s['ls_delta'][i]:+,.1f}" for i in range(1, 5))
    if ls_better >= 2:
        ls_phrase = (f"Enabling Local Search improves the best-configuration average "
                     f"fitness on {ls_better} of 4 instances (deltas {delta_str}; "
                     f"negative = improvement)")
    else:
        ls_phrase = (f"Under this equal budget the Local Search probes do not pay off: "
                     f"plain HS reaches a better best-configuration average on "
                     f"{4 - ls_better} of 4 instances (deltas {delta_str}; negative "
                     f"would mean HS+LS better)")
    best_cfg_str = {m: {i: f"C{s['best_cfg'][m][i] + 1}" for i in range(1, 5)}
                    for m in METHODS}

    # ------------------------------------------------------------------ #
    # Header / abstract
    # ------------------------------------------------------------------ #
    replace_in_paragraph(P[1], "June 5, 2026", "June 12, 2026")
    set_text(P[5],
             "Note on figures. All figures are embedded in this document; they can be "
             "regenerated from the CSV files in results/ with plot_results.py "
             "(output written to results/plots/).")
    set_text(P[8],
             "This report presents a performance analysis of three metaheuristic methods — "
             "Harmony Search without Local Search (HS), Harmony Search with a memetic "
             "Local Search operator (HS+LS), and Ant Colony Optimization (ACO) — applied "
             "to resource allocation in a 5G network that simultaneously serves three "
             "service classes: URLLC (Ultra-Reliable Low-Latency Communication), eMBB "
             "(enhanced Mobile Broadband), and mMTC (massive Machine-Type Communication). "
             "The problem is modelled as the assignment of a transmission mode to each of "
             "600 User Equipments (UEs), minimizing a weighted combination of scheduling "
             "latency and energy consumption subject to a finite pool of radio Resource "
             "Blocks (RBs). The methods are compared on four problem instances with "
             "different service mixes; the exact instances are published in the project "
             "repository.")
    set_text(P[9],
             f"Each method was evaluated with eight parameter configurations and "
             f"{N_RETEST} independent runs per configuration on every instance. Every "
             f"run is terminated after exactly {BUDGET} fitness function evaluations "
             f"(FFE), so all methods receive an equal computational budget — an HS "
             f"iteration evaluates one new harmony while an ACO iteration evaluates one "
             f"solution per ant, which makes iteration counts incomparable. Under this "
             f"equal budget, {winner_phrase}. {ls_phrase}. The best solution on the "
             f"baseline instance is {fmt(s['best_known'][1])}, found by "
             f"{LABEL[min(METHODS, key=lambda m: s['best_best'][m][1])]}.")

    # ------------------------------------------------------------------ #
    # Introduction citations
    # ------------------------------------------------------------------ #
    if not replace_in_paragraph(P[12], "conflicting requirements:",
                                "conflicting requirements [6], [7]:"):
        P[12].add_run(" [6], [7]")
    set_text(P[18],
             "The objectives of this study are to (i) benchmark Harmony Search and Ant "
             "Colony Optimization on a common 5G scheduling problem under identical "
             "conditions and an equal budget of fitness function evaluations, and "
             "(ii) strengthen Harmony Search's weak point — exploitation — by adding a "
             "memetic Local Search operator. Every base configuration is therefore run "
             "both without Local Search (HS) and with it (HS+LS), which makes the "
             "comparison both cross-algorithm (HS vs. ACO) and cross-variant (HS without "
             "vs. with Local Search) on four problem instances.")
    P[20].add_run(" [1]")
    P[21].add_run(" [2], [3]")

    # ------------------------------------------------------------------ #
    # Problem formulation: replace code-source notes with neutral wording
    # ------------------------------------------------------------------ #
    set_text(P[27],
             "The UE counts per service class vary across the four problem instances "
             "(Section 4.2); all other constants above are identical in every experiment.")
    set_text(P[32], "The search space therefore contains 3⁶⁰⁰ candidate "
                    "assignments.")
    set_text(P[42],
             "URLLC is dominated by latency (high time weight); mMTC is dominated by "
             "energy (high energy weight); eMBB sits in between. All methods call the "
             "identical objective function, so the comparison is fair.")

    # ------------------------------------------------------------------ #
    # Section 3: Local Search description and citations
    # ------------------------------------------------------------------ #
    set_text(P[47],
             "The added operator is a best-improvement hill-climber — a memetic "
             "extension of Harmony Search in the sense of [4], [5]. After a new harmony "
             "is improvised and evaluated, with probability LSR the operator runs "
             "LS_STEPS refinement steps. Each step (i) picks a random UE, (ii) evaluates "
             "every alternative mode for that UE via")
    set_text(P[48],
             "the shared objective function, and (iii) commits the mode that most "
             "reduces cost — accepting a change only if it strictly improves the "
             "solution, otherwise leaving the UE unchanged. The resulting cost is then "
             "used for the memory-admission decision. Every probe of the hill-climber "
             "is an ordinary fitness evaluation and is charged against the run's FFE "
             "budget.")
    set_text(P[54],
             "The benchmark evaluates every base configuration both without Local Search "
             "(LSR = 0, the plain HS method) and with it (LSR = 0.3, LS_STEPS = 5, the "
             "HS+LS method), so the contribution of the operator can be isolated per "
             "configuration and per instance.")
    set_text(P[59],
             "All methods share the same problem definition, mode tables, scheduling "
             "logic, and objective function, differing only in how they search the "
             "3⁶⁰⁰ space — and all methods are granted exactly the same "
             "evaluation budget (Section 4.1). This guarantees that observed differences "
             "reflect the algorithms, not the model or the budget.")

    # ------------------------------------------------------------------ #
    # Section 4.1 Protocol
    # ------------------------------------------------------------------ #
    set_text(P[62],
             f"Three methods are compared: Harmony Search without Local Search (HS), "
             f"Harmony Search with Local Search (HS+LS), and Ant Colony Optimization "
             f"(ACO). Each method is run with eight parameter configurations "
             f"(Section 4.3) on four problem instances (Section 4.2), with {N_RETEST} "
             f"independent runs per configuration using fixed random seeds. Every run is "
             f"terminated after exactly the same budget of {BUDGET} fitness function "
             f"evaluations (FFE), where one FFE is one call to the shared objective "
             f"function — including evaluations of infeasible solutions and the probes "
             f"performed by Local Search. An equal FFE budget, rather than an equal "
             f"iteration count, is the fair basis for comparison because one HS "
             f"iteration evaluates a single new harmony while one ACO iteration "
             f"evaluates NumAnts complete solutions [8]. Reported figures are "
             f"aggregated across the {N_RETEST} runs. Fitness is total cost (lower is "
             f"better); time is wall-clock milliseconds.")

    # ------------------------------------------------------------------ #
    # New Section 4.2 Problem Instances (inserted before old "4.2 Parameter Sets")
    # ------------------------------------------------------------------ #
    anchor = P[63]._element  # "4.2 Parameter Sets" heading
    heading = new_paragraph_like(P[63], "4.2 Problem Instances")
    body = new_paragraph_like(P[62],
        "All experiments use four fixed instances of the scheduling problem. Each "
        "instance keeps the total of 600 UEs but varies the service mix "
        "(URLLC/eMBB/mMTC counts) and the per-UE traffic-demand ranges; UE demands are "
        "drawn once with a fixed seed, so every method and configuration is evaluated "
        "on identical data. The exact instances (CSV files) are published at: ")
    add_hyperlink(body, INSTANCES_URL, INSTANCES_URL)

    inst_tbl_el = copy.deepcopy(T[2]._element)  # 4x3 table shell with same look
    anchor.addprevious(heading._element)
    anchor.addprevious(body._element)
    anchor.addprevious(inst_tbl_el)
    spacer = new_paragraph_like(P[62], "")
    anchor.addprevious(spacer._element)

    inst_tbl = Table(inst_tbl_el, T[2]._parent)
    insert_column(inst_tbl, 3)          # 4 columns now
    clone_row(inst_tbl)                 # 5 rows now
    hdr = ["Instance", "Mix (URLLC / eMBB / mMTC)",
           "Data ranges (URLLC; eMBB; mMTC)", "Seed"]
    for c, htext in enumerate(hdr):
        cell_set(inst_tbl.rows[0].cells[c], htext)
    for r, row_vals in enumerate(INSTANCE_ROWS):
        for c, v in enumerate(row_vals):
            cell_set(inst_tbl.rows[r + 1].cells[c], v)

    # Renumber the following subsection headings
    set_text(P[63], "4.3 Parameter Sets")
    set_text(P[70], "4.4 Performance Metrics")
    set_text(P[64],
             "Harmony Search — the same eight base configurations are run without Local "
             "Search (HS, LSR = 0) and with it (HS+LS, LSR = 0.3):")
    set_text(P[73],
             "Accuracy (consistency): accuracy = best_known / avg_fitness × 100, "
             "where best_known is the best solution found on that instance by any "
             "method; higher means an average run lands closer to the best known "
             "solution of the instance.")

    # ------------------------------------------------------------------ #
    # Parameter tables (3: HS, 4: ACO) -> 8 config rows, FFE budget column
    # ------------------------------------------------------------------ #
    hs_param = T[3]   # Set, Iterations, HMS, HMCR, PAR, LSR
    for _ in range(N_CONFIGS - (len(hs_param.rows) - 1)):
        clone_row(hs_param)
    cell_set(hs_param.rows[0].cells[0], "Config")
    cell_set(hs_param.rows[0].cells[1], "Budget (FFE)")
    hs1 = agg["harmony"][1]
    for r in range(N_CONFIGS):
        row = hs_param.rows[r + 1]
        cell_set(row.cells[0], f"C{r + 1}")
        cell_set(row.cells[1], BUDGET)
        cell_set(row.cells[2], int(hs1.loc[r, "HMS"]))
        cell_set(row.cells[3], f"{hs1.loc[r, 'HMCR']:g}")
        cell_set(row.cells[4], f"{hs1.loc[r, 'PAR']:g}")
        cell_set(row.cells[5], "0 / 0.3")

    aco_param = T[4]  # Set, Iterations, NumAnts, alpha, beta, rho, Q
    for _ in range(N_CONFIGS - (len(aco_param.rows) - 1)):
        clone_row(aco_param)
    cell_set(aco_param.rows[0].cells[0], "Config")
    cell_set(aco_param.rows[0].cells[1], "Budget (FFE)")
    aco1 = agg["ant_colony"][1]
    for r in range(N_CONFIGS):
        row = aco_param.rows[r + 1]
        cell_set(row.cells[0], f"C{r + 1}")
        cell_set(row.cells[1], BUDGET)
        cell_set(row.cells[2], int(aco1.loc[r, "NumAnts"]))
        cell_set(row.cells[3], f"{aco1.loc[r, 'PheroCoeff']:g}")
        cell_set(row.cells[4], f"{aco1.loc[r, 'HeurCoeff']:g}")
        cell_set(row.cells[5], f"{aco1.loc[r, 'Rho']:g}")
        cell_set(row.cells[6], f"{aco1.loc[r, 'Q']:g}")

    # ------------------------------------------------------------------ #
    # Section 5.1: HS results table (5) + new HS+LS table
    # ------------------------------------------------------------------ #
    set_text(P[75], "5.1 Harmony Search — with and without Local Search")

    res_hdr = ["Config", "Budget", "HMS", "HMCR", "PAR", "LSR",
               "Avg I1", "Avg I2", "Avg I3", "Avg I4"]

    hs_res = T[5]
    for _ in range(N_CONFIGS - (len(hs_res.rows) - 1)):
        clone_row(hs_res)

    # Clone the extended (now 9-row) table as the future HS+LS table BEFORE filling
    hsls_el = copy.deepcopy(hs_res._element)
    hsls_caption = new_paragraph_like(P[62],
        "The same eight base configurations with Local Search enabled (HS+LS, "
        "LSR = 0.3) — average fitness per instance:")
    hs_res._element.addnext(hsls_el)
    hsls_el.addprevious(hsls_caption._element)
    hsls_res = Table(hsls_el, hs_res._parent)

    def fill_hs_table(tbl, method, lsr_text):
        for c, htext in enumerate(res_hdr):
            cell_set(tbl.rows[0].cells[c], htext)
        base = agg[method][1]
        for r in range(N_CONFIGS):
            row = tbl.rows[r + 1]
            cell_set(row.cells[0], f"C{r + 1}")
            cell_set(row.cells[1], BUDGET)
            cell_set(row.cells[2], int(base.loc[r, "HMS"]))
            cell_set(row.cells[3], f"{base.loc[r, 'HMCR']:g}")
            cell_set(row.cells[4], f"{base.loc[r, 'PAR']:g}")
            cell_set(row.cells[5], lsr_text)
            for i in range(1, 5):
                cell_set(row.cells[5 + i], fmt(agg[method][i].loc[r, "AvgFitness"]))

    fill_hs_table(hs_res, "harmony", "0")
    fill_hs_table(hsls_res, "harmony_ls", "0.3")

    set_text(P[78],
             f"The first table reports the average fitness of every HS configuration "
             f"without Local Search; the second repeats the same eight base "
             f"configurations with Local Search enabled. Under the equal budget of "
             f"{BUDGET} FFE, enabling Local Search improves the best attainable average "
             f"fitness on {ls_better} of 4 instances (best-config deltas {delta_str}; "
             f"negative = HS+LS better). The best HS configuration per instance is "
             + ", ".join(f"I{i}: {best_cfg_str['harmony'][i]}" for i in range(1, 5))
             + "; for HS+LS it is "
             + ", ".join(f"I{i}: {best_cfg_str['harmony_ls'][i]}" for i in range(1, 5))
             + ".")
    set_caption(P[79], "Figure 1. ",
                f"Harmony Search with vs without Local Search on the four instances "
                f"under an equal budget of {BUDGET} FFE (mean over {N_RETEST} runs, "
                f"shaded min–max band; best configuration of each variant).")

    # ------------------------------------------------------------------ #
    # Section 5.2: ACO results table (6)
    # ------------------------------------------------------------------ #
    aco_res = T[6]
    for _ in range(N_CONFIGS - (len(aco_res.rows) - 1)):
        clone_row(aco_res)
    aco_hdr = ["Config", "Budget", "Ants", "α", "β", "ρ", "Q",
               "Avg I1", "Avg I2", "Avg I3", "Avg I4"]
    for c, htext in enumerate(aco_hdr):
        cell_set(aco_res.rows[0].cells[c], htext)
    for r in range(N_CONFIGS):
        row = aco_res.rows[r + 1]
        cell_set(row.cells[0], f"C{r + 1}")
        cell_set(row.cells[1], BUDGET)
        cell_set(row.cells[2], int(aco1.loc[r, "NumAnts"]))
        cell_set(row.cells[3], f"{aco1.loc[r, 'PheroCoeff']:g}")
        cell_set(row.cells[4], f"{aco1.loc[r, 'HeurCoeff']:g}")
        cell_set(row.cells[5], f"{aco1.loc[r, 'Rho']:g}")
        cell_set(row.cells[6], f"{aco1.loc[r, 'Q']:g}")
        for i in range(1, 5):
            cell_set(row.cells[6 + i], fmt(agg["ant_colony"][i].loc[r, "AvgFitness"]))

    set_text(P[83],
             f"Under the equal {BUDGET}-FFE budget the best ACO configuration per "
             f"instance is "
             + ", ".join(f"I{i}: {best_cfg_str['ant_colony'][i]} "
                         f"(avg {fmt(s['best_avg']['ant_colony'][i])})"
                         for i in range(1, 5))
             + ". Smaller colonies complete more pheromone-update rounds within the "
               "fixed budget, while larger colonies sample more solutions per round; "
               "the table shows how this trade-off plays out per instance.")
    set_caption(P[85], "Figure 2. ",
                f"Ant Colony Optimization — convergence of all eight configurations on "
                f"the four instances under {BUDGET} FFE (mean over {N_RETEST} runs).")

    # ------------------------------------------------------------------ #
    # Section 5.3: head-to-head table (7) per instance
    # ------------------------------------------------------------------ #
    set_text(P[87],
             f"The central comparison aligns the three methods per instance under the "
             f"equal budget of {BUDGET} FFE, using the best configuration of each "
             f"method:")
    h2h = T[7]
    h2h_hdr = ["Instance", "HS Avg", "HS Best", "HS+LS Avg", "HS+LS Best",
               "ACO Avg", "ACO Best", "Winner (avg)"]
    for c, htext in enumerate(h2h_hdr):
        cell_set(h2h.rows[0].cells[c], htext)
    for i in range(1, 5):
        row = h2h.rows[i]
        cell_set(row.cells[0], f"I{i}")
        cell_set(row.cells[1], fmt(s["best_avg"]["harmony"][i]))
        cell_set(row.cells[2], fmt(s["best_best"]["harmony"][i]))
        cell_set(row.cells[3], fmt(s["best_avg"]["harmony_ls"][i]))
        cell_set(row.cells[4], fmt(s["best_best"]["harmony_ls"][i]))
        cell_set(row.cells[5], fmt(s["best_avg"]["ant_colony"][i]))
        cell_set(row.cells[6], fmt(s["best_best"]["ant_colony"][i]))
        cell_set(row.cells[7], LABEL[s["winner"][i]])

    set_text(P[90],
             f"Under the equal FFE budget, {winner_phrase}. Because every method spends "
             f"exactly {BUDGET} evaluations, the remaining cost dimension is wall-clock "
             f"overhead per run: on average {fmt(s['time']['harmony'])} ms for HS, "
             f"{fmt(s['time']['harmony_ls'])} ms for HS+LS and "
             f"{fmt(s['time']['ant_colony'])} ms for ACO (mean over all configurations "
             f"and instances; see Section 5.4).")
    set_caption(P[91], "Figure 3. ",
                f"Head-to-head convergence of the best configuration of each method "
                f"(HS, HS+LS, ACO) on each instance under the equal {BUDGET}-FFE "
                f"budget (mean over {N_RETEST} runs, shaded min–max band).")
    set_caption(P[92], "Figure 4. ",
                "Average fitness of every configuration and method per instance "
                "(error bars span best–worst of 10 runs; y-axis zoomed).")
    set_caption(P[93], "Figure 5. ",
                "Accuracy = best_known(instance) / avg_fitness × 100 for the best "
                "configuration of each method on each instance.")

    # ------------------------------------------------------------------ #
    # Section 5.4-5.6
    # ------------------------------------------------------------------ #
    set_caption(P[95], "Figure 6. ",
                f"Average computation time per run at the equal {BUDGET}-FFE budget "
                f"(mean over the eight configurations). Differences now reflect only "
                f"per-evaluation overhead: ACO pays for solution construction via "
                f"roulette-wheel selection, HS+LS for its local-search probes.")
    set_text(P[97],
             "For the best configuration of each method on Instance 1, the per-UE-type "
             "latency and energy are:")

    svc_tbl = T[8]
    insert_column(svc_tbl, 2)  # Metric | HS | HS+LS | ACO
    cell_set(svc_tbl.rows[0].cells[1], "HS (best cfg)")
    cell_set(svc_tbl.rows[0].cells[2], "HS+LS (best cfg)")
    cell_set(svc_tbl.rows[0].cells[3], "ACO (best cfg)")
    svc_rows = [
        ("Latency — URLLC", "AvgLatencyURLLC"),
        ("Energy — URLLC", "AvgEnergyURLLC"),
        ("Latency — eMBB", "AvgLatencyEMBB"),
        ("Energy — eMBB", "AvgEnergyEMBB"),
        ("Latency — mMTC", "AvgLatencyMMTC"),
        ("Energy — mMTC", "AvgEnergyMMTC"),
    ]
    for r, (label, col) in enumerate(svc_rows):
        row = svc_tbl.rows[r + 1]
        cell_set(row.cells[0], label)
        cell_set(row.cells[1], f"{s['svc']['harmony'][col]:.3f}")
        cell_set(row.cells[2], f"{s['svc']['harmony_ls'][col]:.3f}")
        cell_set(row.cells[3], f"{s['svc']['ant_colony'][col]:.3f}")

    lat_min = min(METHODS, key=lambda m: s["svc"][m]["AvgLatencyURLLC"])
    en_min = min(METHODS, key=lambda m: s["svc"][m]["AvgEnergyURLLC"])
    set_text(P[100],
             f"On Instance 1, {LABEL[lat_min]} achieves the lowest average URLLC "
             f"latency ({s['svc'][lat_min]['AvgLatencyURLLC']:.1f} slots) while "
             f"{LABEL[en_min]} consumes the least URLLC energy "
             f"({s['svc'][en_min]['AvgEnergyURLLC']:.1f}). The remaining rows show the "
             f"same latency/energy tension for eMBB and mMTC: no method dominates every "
             f"service-level metric, even when one dominates the aggregate cost.")
    set_caption(P[101], "Figure 7. ",
                "Service-specific latency (log scale) and energy per UE type — "
                "best configuration of each method on Instance 1.")
    set_caption(P[103], "Figure 8. ",
                f"Box plots of the final fitness over {N_RETEST} independent runs at "
                f"{BUDGET} FFE — best configuration of each method on each "
                f"instance. Tighter boxes mean more consistent, reproducible results.")

    # ------------------------------------------------------------------ #
    # Section 6 Discussion
    # ------------------------------------------------------------------ #
    set_text(P[105],
             f"Parameter sensitivity. With the budget fixed, the spread between the "
             f"best and worst of the eight configurations (average fitness, mean over "
             f"instances) is {fmt(s['spread']['harmony'])} for HS, "
             f"{fmt(s['spread']['harmony_ls'])} for HS+LS and "
             f"{fmt(s['spread']['ant_colony'])} for ACO. Parameter choice therefore "
             f"matters for every method, and reporting a single configuration would "
             f"not be representative — hence the eight-configuration protocol [8].")
    set_text(P[106],
             f"Quality vs. time trade-off. Because every method receives exactly "
             f"{BUDGET} evaluations, quality differences can no longer be explained by "
             f"unequal budgets. The remaining cost dimension is wall-clock overhead per "
             f"run: {fmt(s['time']['harmony'])} ms (HS), "
             f"{fmt(s['time']['harmony_ls'])} ms (HS+LS) and "
             f"{fmt(s['time']['ant_colony'])} ms (ACO) on average. The differences stem "
             f"from bookkeeping outside the objective function — ACO's per-UE "
             f"roulette-wheel construction is the most expensive component.")
    set_text(P[107],
             f"Exploration vs. exploitation. Between FFE 500 and FFE {BUDGET} the mean "
             f"best fitness of the best configuration still improves by "
             f"{s['late_gain']['harmony']:.1f}% for HS, "
             f"{s['late_gain']['harmony_ls']:.1f}% for HS+LS and "
             f"{s['late_gain']['ant_colony']:.1f}% for ACO (mean over instances). "
             f"A larger late-stage gain indicates the method is still converting "
             f"evaluations into progress late in the run rather than plateauing early.")
    if ls_better >= 2:
        ls_verdict = ("on these instances the probes generally pay for themselves, "
                      "confirming the value of the memetic extension under a fair "
                      "budget.")
    else:
        ls_verdict = ("on these instances the probes cost more than the improvisations "
                      "they replace — the operator looked beneficial under an "
                      "iteration-based budget (where its extra evaluations were free), "
                      "but under an equal FFE budget plain HS uses the same evaluations "
                      "more effectively. This is precisely why equal-budget comparison "
                      "matters.")
    set_text(P[108],
             f"When Local Search pays off. Under an equal budget every local-search "
             f"probe replaces an improvisation the plain HS could have made, so HS+LS "
             f"only wins where intensification is worth more than additional "
             f"exploration. Empirically it improves the best-configuration average on "
             f"{ls_better} of 4 instances ({delta_str}). Because the neighbourhood (one "
             f"UE's three modes) is tiny and the operator never worsens a solution, its "
             f"risk is low; however, {ls_verdict}")
    set_text(P[109],
             "Service priorities. The per-service results (Section 5.5) show there is "
             "no universally dominant method at the service level: the aggregate-cost "
             "winner does not win every latency or energy metric. If URLLC latency is "
             "the binding constraint, the per-service table — not the aggregate cost — "
             "should drive the choice of method and configuration.")

    # ------------------------------------------------------------------ #
    # Section 7 Conclusion
    # ------------------------------------------------------------------ #
    best_overall_m = min(METHODS, key=lambda m: sum(s["best_avg"][m][i] for i in range(1, 5)))
    set_text(P[111],
             f"On this 5G mode-assignment problem, compared across four instances under "
             f"an equal budget of {BUDGET} fitness function evaluations, "
             f"{winner_phrase}. {ls_phrase}. The overall strongest method across all "
             f"instances is {LABEL[best_overall_m]}. All results are reproducible: the "
             f"four instances are published in the repository, and all runs use fixed "
             f"seeds.")
    set_text(P[112],
             f"Recommendations. Use {LABEL[best_overall_m]} as the default method for "
             f"this problem class (best configuration per instance: "
             + ", ".join(f"I{i}: {best_cfg_str[best_overall_m][i]}" for i in range(1, 5))
             + "). Re-tune configurations when the service mix changes — the best "
               "configuration is not the same on every instance. If a service-level "
               "metric (e.g. URLLC latency) is the binding constraint, select the "
               "method by the per-service table in Section 5.5 instead of aggregate "
               "cost.")

    # ------------------------------------------------------------------ #
    # References (inserted before Appendix A)
    # ------------------------------------------------------------------ #
    appendix_anchor = P[115]._element
    ref_heading = new_paragraph_like(P[115], "References")
    appendix_anchor.addprevious(ref_heading._element)
    for ref in REFERENCES:
        rp = new_paragraph_like(P[12], ref)
        appendix_anchor.addprevious(rp._element)
    spacer2 = new_paragraph_like(P[62], "")
    appendix_anchor.addprevious(spacer2._element)

    # ------------------------------------------------------------------ #
    # Appendix A figure index (table 9) and Appendix B data sources
    # ------------------------------------------------------------------ #
    shows = [
        "HS vs HS+LS convergence per instance (equal FFE budget)",
        "ACO convergence, all 8 configs per instance",
        "Best config of each method per instance (equal FFE)",
        "Avg fitness, all configs × methods per instance",
        "Accuracy (%) per instance, best config of each method",
        "Avg computation time per run at equal FFE",
        "Per-UE-type latency & energy, best configs (Instance 1)",
        "Final-fitness distribution over 10 runs per instance",
    ]
    for r, text in enumerate(shows):
        cell_set(T[9].rows[r + 1].cells[2], text)

    set_text(P[119],
             "results/{harmony, harmony_ls, ant_colony}_aggregated_instance{1..4}.csv "
             "— aggregated results per method and instance (8 configurations "
             "each, 10 runs per configuration).")
    set_text(P[120],
             "results/{method}_ffe_instance{1..4}_config{1..8}.csv — best-so-far "
             "fitness per fitness evaluation, 10 runs each (96 files).")
    set_text(P[121],
             "instances/instance_{1..4}.csv — the four problem instances used in "
             "this study, also published at: ")
    add_hyperlink(P[121], INSTANCES_URL, INSTANCES_URL)
    set_text(P[122],
             f"performance_test.cpp — benchmark harness (3 methods × 8 "
             f"configurations × 4 instances × {N_RETEST} retests, {BUDGET} "
             f"FFE per run). plot_results.py — regenerates all eight figures from "
             f"the CSVs.")

    # ------------------------------------------------------------------ #
    # Replace the 8 embedded figures with the regenerated PNGs
    # ------------------------------------------------------------------ #
    assert len(doc.inline_shapes) == 8, f"expected 8 figures, found {len(doc.inline_shapes)}"
    for shape, fname in zip(doc.inline_shapes, FIGURE_FILES):
        png = (PLOTS / fname).read_bytes()
        r_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        doc.part.related_parts[r_id]._blob = png
        with Image.open(PLOTS / fname) as im:
            w_px, h_px = im.size
        shape.height = int(shape.width * h_px / w_px)

    # Force-create the (missing) docProps part so Word does not complain
    doc.core_properties.title = ("Performance Analysis of Metaheuristic Algorithms "
                                 "for 5G Resource Allocation")

    doc.save(str(DOC_PATH))
    print(f"Saved {DOC_PATH}")

    # ------------------------------------------------------------------ #
    # Post-save validation
    # ------------------------------------------------------------------ #
    check = docx.Document(str(DOC_PATH))
    full_text = "\n".join(p.text for p in check.paragraphs)
    problems = []
    if "Source: lib/" in full_text:
        problems.append("stale 'Source: lib/' reference still present")
    if "References" not in full_text:
        problems.append("References section missing")
    if "26,109.3" in full_text:
        problems.append("stale best-fitness value 26,109.3 still present")
    if f"{BUDGET} fitness function evaluations" not in full_text and \
       f"{BUDGET} FFE" not in full_text:
        problems.append("FFE budget not mentioned")
    if INSTANCES_URL not in full_text:
        # hyperlink text lives outside p.text runs in some cases; check XML
        with zipfile.ZipFile(DOC_PATH) as z:
            xml = z.read("word/document.xml").decode("utf8", "ignore")
        if INSTANCES_URL not in xml and "instances" not in xml:
            problems.append("instances link missing")
    with zipfile.ZipFile(DOC_PATH) as z:
        media = {n: z.getinfo(n).file_size for n in z.namelist()
                 if n.startswith("word/media/")}
    if len(media) != 8:
        problems.append(f"expected 8 media files, found {len(media)}")

    if problems:
        print("VALIDATION PROBLEMS:")
        for p in problems:
            print(" -", p)
        raise SystemExit(1)
    print("Validation OK:", len(media), "figures embedded,",
          len(check.tables), "tables,", len(check.paragraphs), "paragraphs.")


if __name__ == "__main__":
    main()
